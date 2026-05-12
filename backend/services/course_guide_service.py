# backend/services/course_guide_service.py

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Any

from fastapi import UploadFile
from sqlalchemy.orm import Session

from models.course import Course
from models.course_execution import WeeklyPlan


UPLOAD_ROOT = Path("uploads") / "course_guides"


def utcnow():
    return datetime.now(timezone.utc)


def _safe_filename(name: str) -> str:
    name = name.replace("\\", "/").split("/")[-1]
    return re.sub(r"[^a-zA-Z0-9._-]", "_", name)


def clean_text(text: Optional[str]) -> str:
    if not text:
        return ""

    text = str(text).replace("\x00", "")
    text = text.replace("\u2022", "•")
    text = text.replace("￾", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def save_upload(course_id: str, file: UploadFile) -> str:
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

    ts = utcnow().strftime("%Y%m%d_%H%M%S")
    fname = f"{course_id}_{ts}_{_safe_filename(file.filename or 'course_guide')}"
    fpath = UPLOAD_ROOT / fname

    file.file.seek(0)

    with open(fpath, "wb") as f:
        f.write(file.file.read())

    return str(fpath)


# ======================================================
# TEXT EXTRACTION
# ======================================================

def _extract_pdf_lines_with_layout(path: Path) -> List[Dict[str, Any]]:
    """
    Extract PDF text as line records with x/y positions.

    This is important for course outlines because week numbers and day numbers
    can both look like 01, 02, 03. The week column is usually the left-most
    numeric column, so layout helps us separate Week from Day.
    """
    try:
        import fitz  # PyMuPDF
    except Exception as e:
        print(f"[COURSE_GUIDE] PyMuPDF not installed/available: {e}", flush=True)
        return []

    try:
        doc = fitz.open(str(path))
        records: List[Dict[str, Any]] = []

        for page_index, page in enumerate(doc):
            words = page.get_text("words", sort=True) or []

            # word tuple:
            # x0, y0, x1, y1, word, block_no, line_no, word_no
            grouped: Dict[tuple, List[Any]] = {}

            for w in words:
                if len(w) < 8:
                    continue

                key = (page_index, int(w[5]), int(w[6]))
                grouped.setdefault(key, []).append(w)

            for key, line_words in grouped.items():
                line_words = sorted(line_words, key=lambda x: x[0])

                text = " ".join(str(w[4]) for w in line_words).strip()
                text = clean_text(text)

                if not text:
                    continue

                x0 = min(float(w[0]) for w in line_words)
                y0 = min(float(w[1]) for w in line_words)
                x1 = max(float(w[2]) for w in line_words)
                y1 = max(float(w[3]) for w in line_words)

                records.append(
                    {
                        "page": page_index,
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "text": text,
                    }
                )

        doc.close()

        records.sort(key=lambda r: (r["page"], r["y0"], r["x0"]))
        return records

    except Exception as e:
        print(f"[COURSE_GUIDE] PDF layout extraction failed: {e}", flush=True)
        return []


def _extract_pdf_with_pymupdf(path: Path) -> str:
    records = _extract_pdf_lines_with_layout(path)

    if not records:
        return ""

    return clean_text("\n".join(r["text"] for r in records))


def _extract_pdf_with_pypdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = []

        for page in reader.pages:
            parts.append(page.extract_text() or "")

        return clean_text("\n".join(parts))

    except Exception as e:
        print(f"[COURSE_GUIDE] pypdf extraction failed: {e}", flush=True)
        return ""


def _extract_docx(path: Path) -> str:
    try:
        import docx

        d = docx.Document(str(path))
        parts = []

        for p in d.paragraphs:
            if p.text:
                parts.append(p.text)

        # Course guides often keep weekly plans inside tables.
        for table in d.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    val = clean_text(cell.text)

                    if val:
                        row_text.append(val)

                if row_text:
                    parts.append(" | ".join(row_text))

        return clean_text("\n".join(parts))

    except Exception as e:
        print(f"[COURSE_GUIDE] DOCX extraction failed: {e}", flush=True)
        return ""


def extract_text_best_effort(file_path: str) -> str:
    """
    Extract text from PDF/DOCX without hardcoding course content.

    PDF:
    1. Try PyMuPDF layout extraction first.
    2. If weak/empty, try pypdf.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    print(f"[COURSE_GUIDE] Extracting text from: {path} | ext={ext}", flush=True)

    if not path.exists():
        print("[COURSE_GUIDE] File does not exist after upload.", flush=True)
        return ""

    if ext == ".docx":
        text = _extract_docx(path)
        print(f"[COURSE_GUIDE] DOCX extracted chars={len(text)}", flush=True)
        return text

    if ext == ".pdf":
        text = _extract_pdf_with_pymupdf(path)

        if len(text) < 50:
            print("[COURSE_GUIDE] PyMuPDF returned little/no text. Trying pypdf.", flush=True)
            text = _extract_pdf_with_pypdf(path)

        print(f"[COURSE_GUIDE] PDF extracted chars={len(text)}", flush=True)
        return text

    print("[COURSE_GUIDE] Unsupported file type.", flush=True)
    return ""


# ======================================================
# WEEK PARSING - NO COURSE-SPECIFIC HARDCODING
# ======================================================

def _is_stop_section(line: str) -> bool:
    s = line.lower().strip()

    stop_markers = [
        "general grading policy",
        "grading and general course policies",
        "books / reference materials",
        "reference books",
        "policy:",
    ]

    return any(marker in s for marker in stop_markers)


def _is_outline_marker(line: str) -> bool:
    s = line.lower().strip()

    markers = [
        "course outline",
        "week breakdown",
        "topics covered",
        "weekly plan",
        "lecture plan",
        "teaching plan",
        "course schedule",
    ]

    return any(marker in s for marker in markers)


def _parse_week_number_from_line(line: str) -> Optional[int]:
    """
    Supports:
    01
    1
    Week 1
    Week 01
    01 Topic Name
    Week 01 Topic Name
    """
    s = clean_text(line)
    s = s.replace(":", " ").replace("-", " ").replace("–", " ")

    m = re.match(r"^(?:week\s*)?0?([1-9]|1[0-6])(?:\b|\s|$)", s, flags=re.IGNORECASE)

    if not m:
        return None

    try:
        val = int(m.group(1))
    except Exception:
        return None

    if 1 <= val <= 16:
        return val

    return None


def _strip_leading_week_number(line: str) -> str:
    s = clean_text(line)

    s = re.sub(
        r"^(?:week\s*)?0?([1-9]|1[0-6])[\s:\-–.]*",
        "",
        s,
        flags=re.IGNORECASE,
    )

    return clean_text(s)


def _slice_outline_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Keep only the course-outline area if possible.
    """
    if not records:
        return []

    start = 0
    end = len(records)

    for i, r in enumerate(records):
        if _is_outline_marker(r["text"]):
            start = i
            break

    for j in range(start + 1, len(records)):
        if _is_stop_section(records[j]["text"]):
            end = j
            break

    return records[start:end]


def _extract_week_sections_from_pdf_layout(file_path: str) -> Dict[int, str]:
    """
    Generic PDF layout parser.

    Main idea:
    - In tables, Week column is usually the left-most numeric column.
    - Day columns may also contain 01, 02, 03, but they appear further right.
    - So we detect candidate week numbers using x-position, not course-specific titles.
    """
    path = Path(file_path)

    if path.suffix.lower() != ".pdf":
        return {}

    records = _extract_pdf_lines_with_layout(path)
    records = _slice_outline_records(records)

    if not records:
        return {}

    numeric_records = []

    for idx, r in enumerate(records):
        week_no = _parse_week_number_from_line(r["text"])

        if week_no is not None:
            numeric_records.append((idx, r, week_no))

    if not numeric_records:
        return {}

    # Find the left-most numeric column. This is usually the Week column.
    min_x = min(r["x0"] for _, r, _ in numeric_records)

    # Allow slight PDF extraction differences.
    week_x_limit = min_x + 35

    candidates = []
    last_week = 0

    for idx, r, week_no in numeric_records:
        # Must be in the left-most numeric area to avoid day numbers.
        if r["x0"] > week_x_limit:
            continue

        # Course schedules should move forward from week 1 to 16.
        # This avoids treating repeated day numbers as new weeks.
        if week_no <= last_week:
            continue

        # Avoid huge accidental jumps unless the document really skipped weeks.
        if week_no > 16:
            continue

        candidates.append((idx, week_no))
        last_week = week_no

    if not candidates:
        return {}

    sections: Dict[int, str] = {}

    for pos, (start_idx, week_no) in enumerate(candidates):
        end_idx = candidates[pos + 1][0] if pos + 1 < len(candidates) else len(records)

        chunk_lines = []

        for r in records[start_idx:end_idx]:
            line = clean_text(r["text"])

            if not line:
                continue

            if _is_stop_section(line):
                break

            # For the first line, remove just the week number if it has a title.
            if r is records[start_idx]:
                stripped = _strip_leading_week_number(line)
                if stripped:
                    chunk_lines.append(stripped)
                else:
                    chunk_lines.append(f"Week {week_no}")
            else:
                chunk_lines.append(line)

        chunk = clean_text("\n".join(chunk_lines))

        if chunk and len(chunk) > 5:
            sections[week_no] = chunk[:5000]

    return sections


def _extract_week_sections_from_text(text: str) -> Dict[int, str]:
    """
    Generic text parser for DOCX or simple text-based PDFs.

    It does not use course-specific topic names.
    """
    text = clean_text(text)

    if not text:
        return {}

    lines = [clean_text(x) for x in text.splitlines() if clean_text(x)]

    # Keep only outline/schedule area if possible.
    start = 0
    end = len(lines)

    for i, line in enumerate(lines):
        if _is_outline_marker(line):
            start = i
            break

    for j in range(start + 1, len(lines)):
        if _is_stop_section(lines[j]):
            end = j
            break

    lines = lines[start:end]

    candidates = []
    last_week = 0

    for idx, line in enumerate(lines):
        week_no = _parse_week_number_from_line(line)

        if week_no is None:
            continue

        if week_no <= last_week:
            continue

        # Accept sequential or skipped weeks.
        if 1 <= week_no <= 16:
            candidates.append((idx, week_no))
            last_week = week_no

    sections: Dict[int, str] = {}

    for pos, (start_idx, week_no) in enumerate(candidates):
        end_idx = candidates[pos + 1][0] if pos + 1 < len(candidates) else len(lines)

        chunk_lines = []

        for line in lines[start_idx:end_idx]:
            if _is_stop_section(line):
                break

            if not chunk_lines:
                stripped = _strip_leading_week_number(line)
                chunk_lines.append(stripped or f"Week {week_no}")
            else:
                chunk_lines.append(line)

        chunk = clean_text("\n".join(chunk_lines))

        if chunk and len(chunk) > 5:
            sections[week_no] = chunk[:5000]

    return sections


def _extract_assessments_from_week_text(text: str) -> str:
    """
    Pulls assessment/activity hints from the week section dynamically.
    """
    if not text:
        return ""

    patterns = [
        r"\bQuiz\s*\d*\b",
        r"\bAssignment\s*\d*\b",
        r"\bMid[- ]?Term\b",
        r"\bFinal\s+Exam\b",
        r"\bProject\b",
        r"\bPresentation\b",
        r"\bLab\b",
        r"\bActivity\b",
        r"\bClass\s+Task\b",
    ]

    found = []

    for pat in patterns:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            item = clean_text(m.group(0))
            if item and item.lower() not in [x.lower() for x in found]:
                found.append(item)

    return ", ".join(found[:8])


def _fallback_week_plan(week_no: int) -> str:
    return (
        f"Week {week_no} was not clearly detected from the uploaded course guide.\n\n"
        "Manual review may be required. The parser extracts week-wise plans from "
        "numbered week rows such as Week 1, 01, Week 01, etc."
    )


def ensure_weekly_plans(
    db: Session,
    course_id: str,
    planned_text: str,
    file_path: Optional[str] = None,
):
    """
    Creates/updates 16 WeeklyPlan rows.

    No course-specific hardcoding:
    1. Try PDF layout parser using x/y coordinates.
    2. Try generic text parser.
    3. Missing weeks are marked for manual review instead of fake/repeated content.
    """
    planned_text = clean_text(planned_text)

    week_sections: Dict[int, str] = {}

    if file_path:
        week_sections = _extract_week_sections_from_pdf_layout(file_path)

    if len(week_sections) < 4:
        text_sections = _extract_week_sections_from_text(planned_text)

        # Merge without overwriting better PDF-layout results.
        for w, val in text_sections.items():
            if w not in week_sections or len(val) > len(week_sections[w]):
                week_sections[w] = val

    print(
        f"[COURSE_GUIDE] Creating weekly plans | extracted_chars={len(planned_text)} | parsed_weeks={len(week_sections)}",
        flush=True,
    )

    db.query(WeeklyPlan).filter(WeeklyPlan.course_id == course_id).delete()

    now = utcnow()

    for w in range(1, 17):
        topics = week_sections.get(w)

        if topics:
            planned_topics = topics[:5000]
            planned_assessments = _extract_assessments_from_week_text(topics)
        else:
            planned_topics = _fallback_week_plan(w)
            planned_assessments = ""

        wp = WeeklyPlan(
            course_id=course_id,
            week_number=w,
            planned_topics=planned_topics,
            planned_assessments=planned_assessments,
            planned_start_date=None,
            planned_end_date=None,
            created_at=now,
            updated_at=now,
        )

        db.add(wp)

    db.commit()


def set_course_guide_metadata(
    db: Session,
    course: Course,
    file_path: str,
    extracted_text: str,
):
    if hasattr(course, "course_guide_path"):
        course.course_guide_path = file_path

    if hasattr(course, "course_guide_text"):
        course.course_guide_text = clean_text(extracted_text)[:20000]

    db.add(course)
    db.commit()